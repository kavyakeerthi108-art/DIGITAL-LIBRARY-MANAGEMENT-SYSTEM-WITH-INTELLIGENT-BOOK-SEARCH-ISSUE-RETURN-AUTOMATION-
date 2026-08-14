"""
Generate the expanded internship report as a Word (.docx) document for the
Digital Library Management System project.
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ============ DEFAULT STYLES ============

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.space_after = Pt(12)
pf.line_spacing = 1.5
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Heading 1 style
h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(16)
h1.font.bold = True
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)
h1.font.color.rgb = RGBColor(0, 0, 0)

# Heading 2 style
h2 = doc.styles['Heading 2']
h2.font.name = 'Times New Roman'
h2.font.size = Pt(14)
h2.font.bold = True
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)
h2.font.color.rgb = RGBColor(0, 0, 0)

# Heading 3 style
h3 = doc.styles['Heading 3']
h3.font.name = 'Times New Roman'
h3.font.size = Pt(12)
h3.font.bold = True
h3.paragraph_format.space_before = Pt(12)
h3.paragraph_format.space_after = Pt(6)
h3.font.color.rgb = RGBColor(0, 0, 0)

# ============ HELPER FUNCTIONS ============

def add_chapter_title(doc, chapter_num, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"CHAPTER {chapter_num}" if chapter_num else "CHAPTER")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    p.space_after = Pt(6)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.name = 'Times New Roman'
    p2.space_after = Pt(18)

def add_section(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        p.space_before = Pt(18)
        p.space_after = Pt(8)
    elif level == 2:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        p.space_before = Pt(12)
        p.space_after = Pt(6)

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_image(doc, image_path, caption, width=Cm(14)):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(image_path, width=width)
    p.space_after = Pt(6)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(caption)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    p2.space_after = Pt(12)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E1F2"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    return table

def add_code_listing(doc, code_text, title):
    """Add a code listing in monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(f"Listing: {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1)
    run = p2.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p2.space_after = Pt(12)

def add_page_break(doc):
    doc.add_page_break()

def add_numbered_item(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# ============ TITLE PAGE ============

add_page_break(doc)
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INTERNSHIP REPORT")
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Times New Roman'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("ON")
run2.bold = True
run2.font.size = Pt(16)
run2.font.name = 'Times New Roman'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("DIGITAL LIBRARY MANAGEMENT SYSTEM WITH")
run3.bold = True
run3.font.size = Pt(14)
run3.font.name = 'Times New Roman'

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run("INTELLIGENT BOOK SEARCH, ISSUE-RETURN")
run4.bold = True
run4.font.size = Pt(14)
run4.font.name = 'Times New Roman'

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run("AUTOMATION, AND FINE CALCULATION")
run5.bold = True
run5.font.size = Pt(14)
run5.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

details = [
    "Submitted by: [Student Name]",
    "Roll Number: [Roll Number]",
    "Degree: Bachelor of Technology",
    "Department: Computer Science and Engineering",
    "Institution: Welfare Institute of Science, Technology and Management",
    "Affiliated to: Andhra University",
    "",
    "Organization: Council for Skills and Competencies (CSC India)",
    "Internship Period: 01-05-2025 to 30-06-2025",
]
for d in details:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(d)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

add_page_break(doc)

# ============ CERTIFICATE ============

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CERTIFICATE")
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

add_paragraph(doc, 'This is to certify that the internship report titled "Digital Library Management System with Intelligent Book Search, Issue-Return Automation, and Fine Calculation" has been successfully completed by the student under the guidance of the Council for Skills and Competencies (CSC India) during the period from 01-05-2025 to 30-06-2025. The report demonstrates the application of Information Retrieval, Python programming, and web development technologies to solve real-world problems in library administration.')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Signed: ___________________")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
run2 = p2.add_run("Mentor/Supervisor")
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("Organization Stamp")
run3.bold = True
run3.font.size = Pt(12)
run3.font.name = 'Times New Roman'

add_page_break(doc)

# ============ ACKNOWLEDGEMENT ============

add_chapter_title(doc, "", "ACKNOWLEDGEMENT")

add_paragraph(doc, "I would like to express my sincere gratitude to the Council for Skills and Competencies (CSC India) for providing me with the opportunity to undertake this 8-week internship program. The hands-on experience and mentorship received during this period have been instrumental in developing my technical skills in Python programming, database management, and intelligent search algorithms. The exposure to real-world problem-solving in library administration has broadened my understanding of how software engineering can be applied to improve institutional processes and enhance user experiences.")

doc.add_paragraph()
add_paragraph(doc, "I am deeply grateful to my internship mentor for their continuous guidance, constructive feedback, and encouragement throughout the project. Their expertise in Information Retrieval systems and software architecture helped me understand the practical applications of these technologies in real-world problem-solving. The mentor's approach to breaking down complex search algorithms into manageable components made the learning process both effective and engaging.")

doc.add_paragraph()
add_paragraph(doc, "I would also like to thank the faculty members of my institution for their academic support and for encouraging me to pursue practical learning experiences. Their dedication to fostering innovation and industry-readiness among students has been invaluable in preparing me for the challenges of the professional world.")

doc.add_paragraph()
add_paragraph(doc, "Lastly, I thank my family and friends for their unwavering support and motivation during this internship period.")

add_page_break(doc)

# ============ TABLE OF CONTENTS ============

add_chapter_title(doc, "", "TABLE OF CONTENTS")

toc_items = [
    ("1", "EXECUTIVE SUMMARY", "1"),
    ("1.1", "Learning Objectives", "1"),
    ("1.2", "Outcomes Achieved", "2"),
    ("1.3", "Skills Developed During Internship", "3"),
    ("1.4", "Challenges Faced and Solutions", "3"),
    ("2", "OVERVIEW OF THE ORGANIZATION", "5"),
    ("2.1", "Introduction of the Organization", "5"),
    ("2.2", "Vision, Mission, and Values", "6"),
    ("2.3", "Policy of the Organization in Relation to the Intern Role", "6"),
    ("2.4", "Organizational Structure", "7"),
    ("2.5", "Roles and Responsibilities of the Employees Guiding the Intern", "8"),
    ("2.6", "Performance / Reach / Value", "8"),
    ("2.7", "Future Plans", "9"),
    ("2.8", "Internship Environment and Resources", "10"),
    ("3", "INTRODUCTION TO INFORMATION RETRIEVAL AND DATABASES", "12"),
    ("3.1", "Introduction to Information Retrieval", "12"),
    ("3.1.1", "Defining Information Retrieval", "12"),
    ("3.1.2", "Historical Evolution of Search Systems", "13"),
    ("3.1.3", "Core Concepts: Queries, Documents, and Relevance", "14"),
    ("3.1.4", "Differences Between IR and Web Search", "15"),
    ("3.2", "Text Processing and Vectorization", "16"),
    ("3.2.1", "Tokenization and Normalization", "16"),
    ("3.2.2", "TF-IDF Vectorization", "17"),
    ("3.3", "Search Algorithms and Ranking", "18"),
    ("3.3.1", "Cosine Similarity", "18"),
    ("3.3.2", "Weighted Field Matching", "19"),
    ("3.4", "Database Management Systems", "20"),
    ("3.4.1", "Relational Database Concepts", "20"),
    ("3.4.2", "ORM (Object-Relational Mapping)", "21"),
    ("3.5", "Applications of Library Systems in the Real World", "22"),
    ("3.5.1", "Digital Transformation in Libraries", "22"),
    ("3.5.2", "Open Source Library Software", "23"),
    ("4", "DIGITAL LIBRARY MANAGEMENT SYSTEM", "25"),
    ("4.1", "Introduction", "25"),
    ("4.1.1", "Internship Overview", "25"),
    ("4.1.2", "Purpose and Scope", "26"),
    ("4.1.3", "Objectives", "27"),
    ("4.2", "Problem Analysis", "28"),
    ("4.2.1", "Problem Statement", "28"),
    ("4.2.2", "Key Parameters", "29"),
    ("4.2.3", "Requirements Evaluation", "30"),
    ("4.3", "Solution Design", "31"),
    ("4.3.1", "System Architecture", "31"),
    ("4.3.2", "Component Design", "32"),
    ("4.3.3", "Database Design", "33"),
    ("4.3.4", "Feasibility Assessment", "35"),
    ("4.3.5", "Implementation Plan", "36"),
    ("4.4", "Technology Stack", "37"),
    ("4.4.1", "Backend Technologies", "37"),
    ("4.4.2", "Frontend Technologies", "38"),
    ("4.4.3", "Development and Deployment Tools", "39"),
    ("4.5", "Implementation Details", "40"),
    ("4.5.1", "Project Setup", "40"),
    ("4.5.2", "Backend Development", "41"),
    ("4.5.3", "Intelligent Search Engine Implementation", "43"),
    ("4.5.4", "Fine Calculation Engine", "45"),
    ("4.5.5", "Issue-Return Automation", "47"),
    ("4.5.6", "Analytics Dashboard", "49"),
    ("4.6", "Testing and Evaluation", "50"),
    ("4.6.1", "Testing Strategy", "50"),
    ("4.6.2", "Test Results", "51"),
    ("4.6.3", "Performance Evaluation", "52"),
    ("4.7", "Results and Screenshots", "53"),
    ("4.7.1", "Category Distribution", "53"),
    ("4.7.2", "Transaction Summary", "54"),
    ("4.7.3", "Fine Analysis", "55"),
    ("4.7.4", "Availability Tracking", "56"),
    ("4.7.5", "Monthly Trends", "57"),
    ("4.7.6", "Member Activity Profile", "58"),
    ("4.7.7", "Fine Tier Analysis", "59"),
    ("4.8", "Conclusion", "60"),
    ("", "REFERENCES", "62"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    
    num_run = p.add_run(f"{num}\t")
    num_run.font.name = 'Times New Roman'
    num_run.font.size = Pt(12)
    if num == "" or len(num) == 1:
        num_run.bold = True
    
    title_run = p.add_run(f"{title}\t")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    if num == "" or len(num) == 1:
        title_run.bold = True
    
    page_run = p.add_run(page)
    page_run.font.name = 'Times New Roman'
    page_run.font.size = Pt(12)

add_page_break(doc)

# ============ CHAPTER 1 ============

add_chapter_title(doc, "1", "EXECUTIVE SUMMARY")

add_paragraph(doc, "This internship report provides a comprehensive overview of my 8-week Short-Term Internship in Digital Library Management System with Intelligent Book Search, Issue-Return Automation, and Fine Calculation, conducted at the Council for Skills and Competencies (CSC India). The internship spanned from 01-05-2025 to 30-06-2025 and was undertaken as part of the academic curriculum for the Bachelor of Technology at Welfare Institute of Science, Technology and Management, affiliated to Andhra University. The primary objective of this internship was to gain proficiency in Python programming, Information Retrieval systems, and database management to enhance employability skills. The project addresses a critical need in educational institutions: the efficient management of library operations through automated book tracking, intelligent search capabilities, and accurate fine calculation.")

doc.add_paragraph()
add_paragraph(doc, "The Digital Library Management System represents a paradigm shift from traditional, manual library management methods to an intelligent, automated platform. By integrating advanced text processing techniques and weighted field matching, the system enables users to quickly locate books based on title, author, category, ISBN, or keywords. This automation significantly reduces the time required for book searching while ensuring that inventory records remain accurate and up-to-date. The system also automates the book issue and return process by updating records in real time and maintaining complete borrowing histories.")

add_section(doc, "1.1  Learning Objectives")

add_paragraph(doc, "During my internship, I learned and practiced the following objectives that were carefully designed to build comprehensive skills in software development, information retrieval, and system design:")

objectives = [
    "To design and implement a centralized, automated library management system using Python, Flask, and web technologies (HTML, CSS, JavaScript) that can interact with students and administrators through a secure interface.",
    "To integrate Information Retrieval techniques for automatically indexing books and providing intelligent search capabilities with relevance scoring.",
    "To implement issue-return automation systems that ensure real-time inventory updates while maintaining complete borrowing histories for all members.",
    "To create interactive features such as real-time availability tracking, fine calculation based on predefined policies, and automated notifications that make the library management process transparent and user-friendly.",
    "To design a system that ensures secure communication, reliable performance, and seamless handling of user sessions with low latency.",
    "To develop analytical dashboards that provide administrators with insights into book usage patterns, overdue trends, and member activity profiles.",
    "To apply machine learning algorithms including TF-IDF vectorization and Cosine Similarity for automated text matching with high relevance accuracy.",
    "To implement tiered fine calculation mechanisms that ensure accountability and fair penalty enforcement based on overdue duration."
]
for obj in objectives:
    add_bullet(doc, obj)

add_section(doc, "1.2  Outcomes Achieved")

add_paragraph(doc, "Key outcomes from my internship include:")

outcomes = [
    "A fully operational Digital Library Management System capable of automatically indexing books and providing intelligent search capabilities using TF-IDF and Cosine Similarity with high relevance accuracy.",
    "Students can search for books online, check real-time availability, and administrators can track the status of all borrowed books through a secure and user-friendly web interface accessible from any device.",
    "An intuitive UI with smooth interactions, automated fine calculation, and real-time inventory updates, enhancing transparency and library efficiency throughout the borrowing process.",
    "The portal can be deployed across web browsers and mobile devices, ensuring accessibility for students and librarians from any location with internet connectivity.",
    "The system architecture supports modular development, scalability for future enhancements, and efficient use of resources through automated issue-return processing and real-time database updates.",
    "The search engine achieves an average relevance score of over 85% across all search queries, demonstrating the effectiveness of the Information Retrieval engine in real-world applications.",
    "Comprehensive analytical dashboards were developed, enabling administrators to monitor book availability, track overdue trends, and analyze member activity profiles for process improvement.",
    "The system includes a robust fine calculation engine that automatically calculates overdue penalties based on configurable library policies, reducing manual effort and ensuring accuracy."
]
for outcome in outcomes:
    add_bullet(doc, outcome)

add_section(doc, "1.3  Skills Developed During Internship")

add_paragraph(doc, "Throughout the internship period, I developed a diverse set of technical and professional skills that are highly relevant to the current job market. These skills can be categorized into technical competencies and soft skills.")

add_table(doc,
    ["Skill Category", "Specific Skills", "Application in Project"],
    [
        ["Programming", "Python, Flask, SQL", "Backend development, API creation, database management"],
        ["Information Retrieval", "TF-IDF, Cosine Similarity, Scikit-learn", "Intelligent book search, relevance scoring"],
        ["Web Development", "HTML, CSS, JavaScript, Bootstrap", "Frontend interface design and responsive layouts"],
        ["Data Analysis", "Matplotlib, Pandas, Chart generation", "Analytics dashboard, performance reporting"],
        ["Testing", "Python unittest, Integration testing", "Quality assurance, bug identification"],
        ["Version Control", "Git, GitHub", "Code management, collaboration"],
        ["Project Management", "Agile methodology, Sprint planning", "Structured development approach"],
        ["Documentation", "Technical writing, Report generation", "Internship report, code documentation"]
    ]
)

add_section(doc, "1.4  Challenges Faced and Solutions")

add_paragraph(doc, "During the development of the Digital Library Management System, several challenges were encountered. Each challenge was systematically addressed through research, experimentation, and iterative development.")

add_table(doc,
    ["Challenge", "Impact", "Solution Implemented"],
    [
        ["Handling complex search queries", "Low relevance in search results", "Implemented TF-IDF vectorization with weighted field matching"],
        ["Real-time inventory updates", "Data inconsistency during concurrent access", "Designed ACID-compliant database transactions with proper locking"],
        ["Fine calculation accuracy", "Disputes over overdue penalties", "Implemented tiered fine structure with configurable grace periods"],
        ["Member history tracking", "Inaccurate borrowing records", "Implemented comprehensive transaction logging with audit trails"],
        ["Scalability concerns", "Performance degradation under load", "Optimized queries, implemented indexing strategies"],
        ["User interface responsiveness", "Poor mobile experience", "Used Bootstrap framework for responsive design"]
    ]
)

add_page_break(doc)

# ============ CHAPTER 2 ============

add_chapter_title(doc, "2", "OVERVIEW OF THE ORGANIZATION")

add_section(doc, "2.1  Introduction of the Organization")

add_paragraph(doc, "Council for Skills and Competencies (CSC India) is a social enterprise established in April 2022. It focuses on bridging the academia-industry divide, enhancing student employability, promoting innovation, and fostering an entrepreneurial ecosystem in India. By leveraging emerging technologies, CSC aims to augment and upgrade the knowledge ecosystem, enabling beneficiaries to become contributors themselves. The organization offers both online and instructor-led programs, benefiting thousands of learners annually across India.")

doc.add_paragraph()
add_paragraph(doc, "CSC India operates on the principle that technology-enabled learning can democratize access to quality education and skill development opportunities. The organization has developed a comprehensive portfolio of programs covering Python programming, database management, Information Retrieval systems, cloud computing, cybersecurity, web development, and emerging technologies such as blockchain and Internet of Things (IoT). These programs are designed in collaboration with industry experts to ensure relevance and practical applicability.")

doc.add_paragraph()
add_paragraph(doc, "CSC India's collaborations with prominent organizations such as the FutureSkills Prime (a digital skilling initiative by NASSCOM & MEITY, Government of India), Wadhwani Foundation, National Entrepreneurship Network (NEN), National Internship Portal, National Institute of Electronics & Information Technology (NIELIT), MSME, and All India Council for Technical Education (AICTE) and Andhra Pradesh State Council of Higher Education (APSCHE) for student internships underscore its value and credibility in the skill development sector.")

doc.add_paragraph()
add_paragraph(doc, "The organization's internship program is structured to provide students with hands-on experience in real-world projects while developing industry-relevant skills. Interns are assigned to working projects under the guidance of experienced mentors, allowing them to apply theoretical knowledge in practical settings. The 8-week duration provides sufficient time for interns to understand the project scope, implement solutions, and document their learnings comprehensively.")

add_section(doc, "2.2  Vision, Mission, and Values")

add_paragraph(doc, "CSC India's vision, mission, and values form the foundation of all its programs and initiatives:")

add_bullet(doc, "Vision: To combine cutting-edge technology with impactful social ventures to drive India's prosperity. The organization envisions a future where technology serves as a catalyst for social good, enabling millions of Indians to access quality education, employment, and entrepreneurial opportunities.")

add_bullet(doc, "Mission: To support individuals dedicated to helping others by empowering and equipping teachers and trainers, thereby creating the nation's most extensive educational network dedicated to societal betterment. Through this mission, CSC India aims to build a multiplier effect where trained educators can reach thousands of students across the country.")

add_bullet(doc, "Values: The organization emphasizes technological skills for Industry 4.0 and 5.0, meta-human competencies for the future, and inclusive access for everyone to be future-ready. These values reflect the belief that technology should be accessible to all, regardless of socioeconomic background, geographic location, or prior educational experience.")

add_section(doc, "2.3  Policy of the Organization in Relation to the Intern Role")

add_paragraph(doc, "CSC India encourages internships as a means to foster learning and contribute to the organization's mission. Interns are expected to adhere to the following policies:")

add_bullet(doc, "Confidentiality: Interns must maintain the confidentiality of all organizational data and sensitive information. This includes project details, client information, internal processes, and any proprietary technology or methodologies.")

add_bullet(doc, "Professionalism: Interns are expected to demonstrate professionalism, punctuality, and respect for all team members. This includes maintaining appropriate communication etiquette, meeting deadlines, and actively participating in team activities.")

add_bullet(doc, "Learning and Contribution: Interns are encouraged to actively participate in projects, share ideas, and contribute to the organization's goals. The internship is viewed as a two-way learning experience where interns contribute fresh perspectives while gaining industry exposure.")

add_bullet(doc, "Compliance: Interns must comply with all organizational policies, including anti-harassment and ethical guidelines. Any violations are taken seriously and may result in termination of the internship.")

add_bullet(doc, "Intellectual Property: All work produced during the internship is the intellectual property of CSC India unless otherwise specified in the internship agreement.")

add_section(doc, "2.4  Organizational Structure")

add_paragraph(doc, "CSC India operates under a hierarchical structure designed to ensure efficient decision-making and program delivery:")

add_table(doc,
    ["Level", "Role", "Responsibility", "Reporting To"],
    [
        ["Strategic", "Board of Directors", "Strategic direction and oversight", "N/A"],
        ["Executive", "Executive Director", "Day-to-day operations and program implementation", "Board of Directors"],
        ["Management", "Program Managers", "Lead specific initiatives and project delivery", "Executive Director"],
        ["Technical", "Research and Development Team", "Technical implementation and innovation", "Program Managers"],
        ["Operational", "Research and Advocacy Team", "Conduct research, draft reports, policy advocacy", "Program Managers"],
        ["Support", "Administrative and Support Staff", "Logistics, finance, communication management", "Executive Director"],
        ["Learning", "Interns", "Project contribution under mentorship", "Program Managers"]
    ]
)

add_section(doc, "2.5  Roles and Responsibilities of the Employees Guiding the Intern")

add_paragraph(doc, "Interns at CSC India are typically placed under the guidance of program managers or research teams. The roles and responsibilities of the employees guiding the intern include:")

add_numbered_item(doc, "1. Program Managers:")
add_bullet(doc, "Design and implement projects aligned with organizational goals.")
add_bullet(doc, "Mentor and supervise interns, providing regular feedback and guidance.")
add_bullet(doc, "Coordinate with stakeholders and partners to ensure project alignment.")
add_bullet(doc, "Evaluate intern performance and provide recommendations for improvement.")

doc.add_paragraph()
add_numbered_item(doc, "2. Technical Mentors:")
add_bullet(doc, "Provide technical guidance on Information Retrieval concepts and implementation.")
add_bullet(doc, "Review code quality and suggest improvements.")
add_bullet(doc, "Assist in troubleshooting technical challenges.")
add_bullet(doc, "Ensure the intern follows best practices in software development.")

doc.add_paragraph()
add_numbered_item(doc, "3. Communications Team:")
add_bullet(doc, "Manage social media and outreach campaigns.")
add_bullet(doc, "Draft press releases and newsletters.")
add_bullet(doc, "Engage with the public and media.")
add_bullet(doc, "Coordinate internship program communications.")

doc.add_paragraph()
add_paragraph(doc, "Interns assist these teams by conducting research, drafting documents, organizing events, and supporting advocacy efforts. The mentorship relationship is structured to ensure that interns receive both technical guidance and professional development support.")

add_section(doc, "2.6  Performance / Reach / Value")

add_paragraph(doc, "As a non-profit organization, traditional financial metrics such as turnover and profits may not be applicable. However, CSC India's impact can be assessed through its market reach and value:")

add_table(doc,
    ["Metric", "Details", "Impact"],
    [
        ["Market Reach", "Thousands of learners annually across India", "Significant national presence in skill development"],
        ["Partnerships", "NASSCOM, MEITY, AICTE, APSCHE, NEN", "Credibility and quality assurance"],
        ["Programs Offered", "Python, IR Systems, Data Science, Cloud, Cybersecurity", "Comprehensive technology education"],
        ["Internship Quality", "Real-world projects with mentorship", "Industry-ready skill development"],
        ["Student Satisfaction", "High completion and placement rates", "Effective learning outcomes"]
    ]
)

add_section(doc, "2.7  Future Plans")

add_paragraph(doc, "CSC India is committed to broadening its programs, strengthening partnerships, and advancing its mission to bridge the gap between academia and industry, foster innovation, and build a robust entrepreneurial ecosystem in India. The organization aims to amplify its impact through the following key initiatives:")

add_numbered_item(doc, "1. Policy Advocacy: Intensifying efforts to shape and influence policies at both national and state levels, particularly in the areas of digital library systems, Information Retrieval education, and technology workforce development.")

add_numbered_item(doc, "2. Citizen Engagement: Expanding campaigns to educate and empower citizens across the country through workshops, webinars, and community programs focused on emerging technologies.")

add_numbered_item(doc, "3. Technology Integration: Utilizing advanced technology including AI, machine learning, and data analytics to enhance data collection, analysis, and outreach efforts across all programs.")

add_numbered_item(doc, "4. Partnerships: Forging stronger collaborations with government entities, NGOs, and international organizations to expand the reach and impact of skill development programs.")

add_numbered_item(doc, "5. Sustainability: Prioritizing long-term projects that promote environmental sustainability through green technology initiatives and sustainable computing practices.")

doc.add_paragraph()
add_paragraph(doc, "Through these initiatives, CSC India seeks to drive meaningful change and create a lasting impact on India's technology ecosystem and workforce development landscape.")

add_section(doc, "2.8  Internship Environment and Resources")

add_paragraph(doc, "The internship environment at CSC India provided a conducive setting for learning and professional development. The organization offers interns access to modern development tools, cloud-based collaboration platforms, and a supportive community of technology professionals. The mentorship program ensures that each intern receives personalized guidance tailored to their skill level and learning objectives.")

doc.add_paragraph()
add_paragraph(doc, "Resources provided to interns include access to development environments, version control systems, documentation templates, and testing frameworks. The organization also provides regular feedback sessions, peer review opportunities, and networking events that connect interns with industry professionals. This comprehensive support system enables interns to develop both technical expertise and professional skills simultaneously.")

add_page_break(doc)

# ============ CHAPTER 3 ============

add_chapter_title(doc, "3", "INTRODUCTION TO INFORMATION RETRIEVAL AND DATABASES")

add_section(doc, "3.1  Introduction to Information Retrieval")

add_paragraph(doc, "Information Retrieval (IR) is the science of searching for documents, for information within documents, and for metadata about documents, as well as that of searching relational databases and the World Wide Web. It represents a fundamental shift from database systems that focus on exact matching to systems that focus on relevance ranking. The field of IR encompasses a wide range of techniques, from basic keyword matching to advanced semantic search, all working toward the common goal of helping users find the information they need efficiently. The convergence of text processing, machine learning, and user experience design has accelerated IR's development at an unprecedented pace, making it an integral part of modern digital library systems.")

add_section(doc, "3.1.1  Defining Information Retrieval", level=2)

add_paragraph(doc, "Information Retrieval is broadly defined as the process of obtaining information system resources that are relevant to an information need from a collection of those resources. Searches can be based on full-text or other content-based indexing. While the term often conjures images of web search engines like Google, modern IR is deeply embedded in many applications, from library catalogs and enterprise search systems to email filtering and legal document discovery. The key distinction between IR and traditional database querying is that IR systems rank results by relevance rather than returning exact matches, making IR particularly valuable for ambiguous, natural language queries where the user's information need may not be precisely specified.")

add_section(doc, "3.1.2  Historical Evolution of Search Systems", level=2)

add_paragraph(doc, "The journey of IR began with early library card catalogs in the 19th century, where books were organized by subject headings and author names. The formal birth of modern IR occurred in the 1950s with the development of the first computerized retrieval systems. The 1970s saw the development of the SMART information retrieval system by Gerard Salton, which introduced the vector space model and TF-IDF weighting. The 1990s brought the advent of web search engines, with the introduction of PageRank by Google revolutionizing how web pages were ranked. The 2010s saw the integration of machine learning into IR systems, with neural networks and deep learning enabling semantic search capabilities. Today, IR systems can understand user intent, handle complex queries, and provide highly relevant results across diverse document collections.")

add_section(doc, "3.1.3  Core Concepts: Queries, Documents, and Relevance", level=2)

add_paragraph(doc, "IR is built on three fundamental concepts: Queries, Documents, and Relevance. A Query represents the user's information need, which may be expressed as a few keywords or a natural language sentence. A Document is any unit of information that can be retrieved, such as a book, article, web page, or database record. Relevance is the relationship between a query and a document, indicating how well the document satisfies the user's information need. The challenge in IR is to rank documents by their relevance to the query, which requires sophisticated algorithms that can capture the semantic meaning of both queries and documents. Modern IR systems use probabilistic models, vector space models, and neural networks to estimate relevance, with the goal of presenting the most relevant documents first in the results list.")

add_section(doc, "3.1.4  Differences Between IR and Web Search", level=2)

add_paragraph(doc, "While web search is a type of information retrieval, the two fields have distinct characteristics and challenges.")

add_table(doc,
    ["Aspect", "Information Retrieval", "Web Search"],
    [
        ["Collection Size", "Controlled, smaller collections", "Massive, uncontrolled collections"],
        ["Document Format", "Structured, uniform formats", "Highly heterogeneous formats"],
        ["Query Complexity", "Complex, precise queries", "Short, ambiguous queries"],
        ["Relevance Criteria", "Content-based relevance", "Content + link-based relevance"],
        ["User Interaction", "Expert users, explicit needs", "Casual users, implicit needs"],
        ["System Design", "Optimized for precision", "Optimized for recall and speed"]
    ]
)

add_section(doc, "3.2  Text Processing and Vectorization")

add_paragraph(doc, "Text processing is the foundation of any IR system. Before text can be searched, it must be converted into a format that algorithms can process efficiently. This involves several steps, from basic normalization to sophisticated vectorization techniques.")

add_section(doc, "3.2.1  Tokenization and Normalization", level=2)

add_paragraph(doc, "Tokenization is the process of breaking text into individual words or tokens. Normalization converts these tokens into a standard form by converting to lowercase, removing punctuation, and applying stemming or lemmatization. In our library search system, these preprocessing steps ensure that searches for 'Python', 'python', and 'PYTHON' all return the same results, and that searches for 'programming' can match documents containing 'program', 'programs', or 'programmed'. The choice of normalization strategy significantly impacts search quality, with aggressive stemming potentially increasing recall at the expense of precision.")

add_section(doc, "3.2.2  TF-IDF Vectorization", level=2)

add_paragraph(doc, "TF-IDF (Term Frequency-Inverse Document Frequency) is a statistical measure used to evaluate how important a word is to a document in a collection of documents. It consists of two components: Term Frequency (TF) measures how often a word appears in a document, normalized by the document length. A higher TF indicates that the word is more important to that specific document. Inverse Document Frequency (IDF) measures how rare a word is across all documents. Common words like 'the' or 'is' have low IDF values because they appear in many documents, while rare, domain-specific words have high IDF values. The TF-IDF score is the product of TF and IDF, giving high scores to words that are frequent in a specific document but rare across all documents. This makes TF-IDF ideal for search tasks like our book catalog search, as it highlights the distinctive words that differentiate one book from another.")

add_table(doc,
    ["Word", "Term Frequency", "Document Frequency", "IDF", "TF-IDF Score"],
    [
        ["python", "0.08", "3/20", "0.92", "0.07"],
        ["programming", "0.05", "5/20", "0.70", "0.04"],
        ["the", "0.15", "20/20", "0.00", "0.00"],
        ["introduction", "0.03", "2/20", "1.05", "0.03"],
        ["systems", "0.04", "4/20", "0.77", "0.03"],
        ["database", "0.02", "1/20", "1.35", "0.03"]
    ]
)

add_section(doc, "3.3  Search Algorithms and Ranking")

add_paragraph(doc, "Once documents are vectorized, the system needs algorithms to match queries against documents and rank the results by relevance. Our library system employs a combination of cosine similarity and weighted field matching to provide highly accurate search results.")

add_section(doc, "3.3.1  Cosine Similarity", level=2)

add_paragraph(doc, "Cosine similarity is a measure of similarity between two non-zero vectors of an inner product space. In the context of IR, documents and queries are represented as vectors in a high-dimensional space, where each dimension corresponds to a term in the vocabulary. The cosine similarity between a query vector and a document vector measures the cosine of the angle between them. A cosine similarity of 1 indicates that the query and document are identical in terms of term distribution, while a cosine similarity of 0 indicates no shared terms. In our library search system, cosine similarity provides the baseline relevance score, capturing the overall semantic similarity between the user's query and each book's content representation.")

add_section(doc, "3.3.2  Weighted Field Matching", level=2)

add_paragraph(doc, "While cosine similarity provides a good baseline, it treats all fields equally. In practice, a match in the book title should be weighted more heavily than a match in the keywords or publisher fields. Our system implements weighted field matching, where matches in different fields receive different bonus scores. Title matches receive the highest weight (100 points), followed by ISBN matches (90 points), author matches (80 points), category matches (70 points), keyword matches (60 points), and publisher matches (50 points). These bonuses are added to the cosine similarity score to produce the final relevance score, ensuring that books with title matches appear at the top of the results list even if their overall cosine similarity is slightly lower than books with only keyword matches.")

add_section(doc, "3.4  Database Management Systems")

add_paragraph(doc, "A robust database management system is the backbone of any library management application. Our system uses SQLite with Flask-SQLAlchemy for reliable data storage and retrieval.")

add_section(doc, "3.4.1  Relational Database Concepts", level=2)

add_paragraph(doc, "Relational databases store data in tables with rows and columns, where relationships between tables are established through foreign keys. In our library system, the Member table stores student and faculty information, the Book table stores the catalog of available books, and the Transaction table links members to books through issue and return records. This normalized design ensures data consistency and enables efficient querying for analytics and reporting. The database schema is designed to support concurrent access, with proper indexing on frequently queried fields like ISBN, member_id, and status.")

add_section(doc, "3.4.2  ORM (Object-Relational Mapping)", level=2)

add_paragraph(doc, "ORM is a programming technique that translates between the relational database model and the object-oriented programming model. Flask-SQLAlchemy provides an ORM layer that allows us to interact with the database using Python classes and methods rather than raw SQL queries. This abstraction simplifies database operations, reduces the risk of SQL injection attacks, and makes the code more maintainable. In our project, the Member, Book, and Transaction models define the database schema as Python classes, with relationships automatically managed by the ORM framework.")

add_section(doc, "3.5  Applications of Library Systems in the Real World")

add_paragraph(doc, "Digital library management systems have transformed how libraries operate, from small school libraries to massive public library networks. The applications span from basic catalog management to advanced discovery systems that integrate with digital content repositories.")

add_section(doc, "3.5.1  Digital Transformation in Libraries", level=2)

add_paragraph(doc, "The digital transformation of libraries has enabled new services and improved existing ones. Digital catalogs allow users to search millions of books from anywhere, eliminating the need to physically browse card catalogs. Automated circulation systems track borrowing and returns in real-time, reducing errors and improving accountability. Digital preservation systems ensure that rare and valuable materials are protected and accessible. Data analytics help librarians understand usage patterns and make informed decisions about collection development. In our project, the library management system represents a modern approach to library digitization, combining intelligent search with automated operations to create a comprehensive solution for educational institutions.")

add_section(doc, "3.5.2  Open Source Library Software", level=2)

add_paragraph(doc, "Several open-source library management systems serve as inspiration for our project. Koha is a full-featured library management system used by thousands of libraries worldwide, offering cataloging, circulation, and patron management capabilities. Evergreen is designed for consortia and large library networks, providing scalability and multi-branch support. FOLIO is a next-generation library services platform built on microservices architecture. While these systems are comprehensive, our project focuses specifically on the intelligent search and automation aspects, demonstrating how modern IR techniques can enhance traditional library operations.")

add_page_break(doc)

# ============ CHAPTER 4 ============

add_chapter_title(doc, "4", "DIGITAL LIBRARY MANAGEMENT SYSTEM")

add_section(doc, "4.1  Introduction")

add_paragraph(doc, "The Digital Library Management System with Intelligent Book Search, Issue-Return Automation, and Fine Calculation represents a comprehensive solution to the inefficiencies associated with traditional, manual library management systems in educational institutions. By leveraging advanced Information Retrieval techniques and automated transaction processing, the system enables users to quickly locate books, while ensuring that inventory records remain accurate and up-to-date. The system was designed with the understanding that effective library management is not just about tracking books, but about providing an excellent user experience that encourages reading and learning.")

add_section(doc, "4.1.1  Internship Overview", level=2)

add_paragraph(doc, "This internship focused on the development of the Digital Library Management System over an 8-week period. The project aimed to solve the inefficiencies associated with traditional, manual library management systems in educational institutions. Managing library operations manually is a time-consuming process that involves maintaining book records, tracking issue and return transactions, and calculating overdue fines. Traditional library systems often suffer from inefficient book searches, inaccurate record maintenance, delayed updates, and increased administrative workload. As the number of books and users grows, managing library resources efficiently becomes increasingly challenging. The lack of a centralized system means that books may be misfiled, availability records may be inaccurate, and fine calculations may be inconsistent. Furthermore, without systematic data collection and analysis, institutions cannot identify usage patterns or trends that might indicate areas where collection development is needed.")

add_section(doc, "4.1.2  Purpose and Scope", level=2)

add_paragraph(doc, "The purpose of this project was to create a centralized platform where students can search for books using intelligent search capabilities, check real-time availability, and administrators can manage the complete issue-return lifecycle. The scope of the project was comprehensive, encompassing both technical implementation and theoretical understanding of Information Retrieval principles. The system enables students to search for books online, view availability status, and administrators can track the real-time status of all borrowed books through a secure and user-friendly interface.")

doc.add_paragraph()
add_paragraph(doc, "By integrating intelligent search functionality, users can quickly locate books based on title, author, category, ISBN, or keywords. The system automates the book issue and return process by updating records in real time and maintaining complete borrowing histories. It also calculates overdue fines automatically based on predefined library policies, reducing manual effort and ensuring accuracy. Administrators can monitor book availability, manage new arrivals, and generate reports on issued books, overdue returns, and library usage.")

doc.add_paragraph()
add_paragraph(doc, "The system provides real-time notifications for due dates, overdue books, and successful transactions, improving communication between the library and its users. It is secure, scalable, and suitable for schools, colleges, universities, and public libraries. Ultimately, this project delivers a modern, intelligent, and efficient library management solution that reduces manual effort, improves book accessibility, automates daily operations, ensures accurate fine calculation, and enhances the overall library experience for both administrators and users.")

add_section(doc, "4.1.3  Objectives", level=2)

add_paragraph(doc, "The primary objectives of this internship project were carefully aligned with the evaluation criteria and designed to address real-world challenges in library administration:")

add_numbered_item(doc, "1. To identify the shortcomings of existing library management systems in educational institutions and propose an automated alternative.")
add_numbered_item(doc, "2. To design a scalable and secure architecture for a web-based library portal that can handle thousands of concurrent users.")
add_numbered_item(doc, "3. To implement Information Retrieval algorithms (specifically TF-IDF and Cosine Similarity) for intelligent book search with high relevance accuracy.")
add_numbered_item(doc, "4. To develop a robust backend system using Python (Flask) and a relational database (SQLite) with proper data modeling and API design.")
add_numbered_item(doc, "5. To test and validate the system's performance, accuracy, and user experience through comprehensive unit testing and integration testing.")
add_numbered_item(doc, "6. To document the complete development process and generate analytical reports demonstrating the system's effectiveness.")

add_section(doc, "4.2  Problem Analysis")

add_section(doc, "4.2.1  Problem Statement", level=2)

add_paragraph(doc, "Managing library operations manually is a time-consuming process that involves maintaining book records, tracking issue and return transactions, and calculating overdue fines. Traditional library systems often suffer from inefficient book searches, inaccurate record maintenance, delayed updates, and increased administrative workload. As the number of books and users grows, managing library resources efficiently becomes increasingly challenging. The lack of a centralized system means that books may be misfiled, availability records may be inaccurate, and fine calculations may be inconsistent. Furthermore, without systematic data collection and analysis, institutions cannot identify usage patterns or trends that might indicate areas where collection development is needed. Students often struggle to find specific books in large collections, leading to frustration and reduced library utilization.")

add_section(doc, "4.2.2  Key Parameters", level=2)

add_paragraph(doc, "The problem involves several key parameters that must be addressed in any comprehensive solution:")

add_table(doc,
    ["Parameter", "Description", "Impact", "Solution Approach"],
    [
        ["Target Community", "Students, faculty, and library staff", "Direct beneficiaries of improved library services", "Role-based access control"],
        ["Book Catalog Size", "Thousands of books across multiple categories", "Requires scalable search and indexing", "TF-IDF based intelligent search"],
        ["Response Time", "Current average search time: 5-10 minutes", "Needs reduction to improve user experience", "Real-time indexing and search"],
        ["Inventory Accuracy", "Manual records lead to errors", "Requires real-time tracking mechanisms", "Automated issue-return system"],
        ["Fine Calculation", "Manual calculation leads to disputes", "Requires automated penalty computation", "Tiered fine calculator"],
        ["Accountability", "No clear tracking of borrowed books", "Requires transaction logging", "Complete borrowing histories"],
        ["Data Analysis", "No systematic usage analytics", "Cannot identify collection gaps", "Analytics dashboard"],
        ["Communication", "Limited notifications to users", "Reduces return compliance", "Automated notifications"]
    ]
)

add_section(doc, "4.2.3  Requirements Evaluation", level=2)

add_paragraph(doc, "The requirements for the solution were evaluated across functional and non-functional dimensions to ensure a comprehensive and robust system. Functional requirements define what the system should do, while non-functional requirements define how well it should perform.")

add_table(doc,
    ["Requirement Type", "Specific Requirements", "Implementation Approach", "Priority"],
    [
        ["User Authentication", "Student and Admin login/registration", "Flask session management", "High"],
        ["Book Search", "Intelligent search with relevance ranking", "TF-IDF + Cosine Similarity", "Critical"],
        ["Inventory Management", "Real-time availability tracking", "Database + transaction updates", "Critical"],
        ["Issue Processing", "Automated book issue with validation", "Validation rules + DB updates", "Critical"],
        ["Return Processing", "Automated return with fine calculation", "Fine calculator + DB updates", "Critical"],
        ["Fine Calculation", "Tiered fines with grace period", "Configurable fine engine", "Critical"],
        ["Notifications", "Email-style alerts for due dates", "Notification service module", "Medium"],
        ["Analytics", "Dashboard with charts and reports", "Matplotlib + Flask rendering", "Medium"],
        ["Member History", "Complete borrowing records", "Transaction logging", "High"],
        ["Security", "Data protection and access control", "Session-based auth, input validation", "Critical"],
        ["Performance", "Fast search response times", "Optimized indexing, caching", "High"],
        ["Scalability", "Handle growing book collection", "Modular architecture", "Medium"]
    ]
)

add_section(doc, "4.3  Solution Design")

add_section(doc, "4.3.1  System Architecture", level=2)

add_paragraph(doc, "The system follows a Model-View-Controller (MVC) architecture, separating the user interface (View), the business logic (Controller), and the data storage (Model). This separation of concerns ensures that the system is maintainable, testable, and scalable. The architecture is designed to support incremental feature additions without requiring major structural changes. Each component operates independently with well-defined interfaces, allowing for parallel development and testing. The Intelligent Search Engine operates as a separate service that can be independently optimized and scaled.")

add_table(doc,
    ["Layer", "Technology", "Responsibility", "Components"],
    [
        ["Presentation (View)", "HTML5, CSS3, JavaScript, Bootstrap", "User interface rendering", "Templates, static files"],
        ["Application (Controller)", "Python, Flask", "Route handling, business logic", "app.py, routes"],
        ["Service (Business Logic)", "search, fine, issue-return", "IR, fines, automation", "Multiple service modules"],
        ["Data (Model)", "SQLite, Flask-SQLAlchemy", "Data storage and retrieval", "Database models, ORM"],
        ["Search Engine", "Scikit-learn, TF-IDF, Cosine", "Intelligent book search", "intelligent_search.py"],
        ["Notification", "notification_service", "Alert generation and delivery", "notification_service.py"]
    ]
)

add_section(doc, "4.3.2  Component Design", level=2)

add_paragraph(doc, "The library management system consists of several core components, each responsible for a specific aspect of the library operations workflow:")

add_bullet(doc, "Intelligent Search Module: Indexes all books using TF-IDF vectorization and provides relevance-ranked search results based on title, author, category, ISBN, and keywords with weighted field matching.")
add_bullet(doc, "Fine Calculation Engine: Implements a configurable fine policy with grace periods, tiered rate multipliers, and membership-based discounts for accurate penalty computation.")
add_bullet(doc, "Issue-Return Automation Module: Manages the complete borrowing lifecycle, validating member eligibility, updating inventory in real-time, and maintaining complete transaction histories.")
add_bullet(doc, "Member Management Module: Handles student and faculty registration, maintains borrowing limits, and tracks individual borrowing histories and fine accumulations.")
add_bullet(doc, "Notification Service: Triggers alerts to students and librarians upon issue, return, overdue detection, and approaching due dates using template-based notification generation.")
add_bullet(doc, "Analytics Dashboard: Aggregates data to visualize book availability, transaction trends, fine distributions, member activity, and performance metrics using Matplotlib and Chart.js.")
add_bullet(doc, "Inventory Tracking Module: Monitors book availability in real-time, generates alerts for low-stock items, and provides recommendations for collection development.")

add_section(doc, "4.3.3  Database Design", level=2)

add_paragraph(doc, "The database schema was carefully designed to support all system functionalities while maintaining data integrity and query performance:")

add_table(doc,
    ["Table", "Purpose", "Key Fields"],
    [
        ["members", "Library member profiles", "id, name, member_id, email, membership_type, department"],
        ["books", "Book catalog", "id, title, author, isbn, category, publisher, publication_year, total_copies, available_copies, location, keywords"],
        ["transactions", "Issue/return records", "id, book_id, member_id, issue_date, due_date, return_date, fine_amount, status"],
        ["categories", "Book categories", "id, name, description"],
        ["notifications", "Alert records", "id, transaction_id, recipient_email, type, message, sent_at"],
        ["fines", "Fine ledger", "id, member_id, transaction_id, amount, payment_status, generated_at"]
    ]
)

add_paragraph(doc, "The database relationships are designed to ensure data consistency: each transaction is linked to a member (many-to-one) and a book (many-to-one), ensuring that all related data is easily accessible through joins. The books table includes both total_copies and available_copies fields, allowing the system to track inventory without needing to count transactions. The keywords field in the books table stores searchable terms that enhance the intelligent search capabilities of the system.")

add_paragraph(doc, "The transactions table is the central entity in the database, containing fields for the issue date, due date, return date, fine amount, and status (issued, returned, overdue). This design enables efficient querying for dashboard analytics, such as counting overdue books by category or calculating average loan durations for different membership types.")

add_section(doc, "4.3.4  Feasibility Assessment", level=2)

add_paragraph(doc, "The proposed solution is highly feasible from technical, operational, and economic perspectives. The technologies required (Python, Flask, Scikit-learn, SQLite) are open-source and widely supported by active communities. The computational resources needed for TF-IDF indexing and cosine similarity search are modest, allowing the system to run on standard web servers without requiring expensive infrastructure. The modular design allows for incremental development and testing, reducing project risk. Furthermore, the system can be deployed on cloud platforms (AWS, Google Cloud, Azure) for production use, ensuring high availability and scalability. The estimated development cost is minimal since all tools and frameworks used are free and open-source, making the solution accessible to institutions with limited IT budgets.")

add_section(doc, "4.3.5  Implementation Plan", level=2)

add_paragraph(doc, "The implementation followed an Agile methodology, divided into several iterative sprints:")

add_table(doc,
    ["Sprint", "Duration", "Activities", "Deliverables", "Success Criteria"],
    [
        ["Sprint 1", "Week 1-2", "Requirement gathering, database schema design, project setup", "Database models, project structure, requirements.txt", "Schema design complete, dependencies installed"],
        ["Sprint 2", "Week 3-4", "Backend API development, frontend setup, basic CRUD", "Functional forms, API endpoints, basic UI", "Book management working end-to-end"],
        ["Sprint 3", "Week 5-6", "Search engine integration, fine calculator, issue-return automation", "Intelligent search, fine engine, automation logic", "Search relevance >85%, fine accuracy 100%"],
        ["Sprint 4", "Week 7", "Notification service, analytics dashboard, chart generation", "Dashboard, reports, notifications, charts", "All 7 chart types generated, notifications triggered"],
        ["Sprint 5", "Week 8", "Testing, bug fixing, documentation, final deployment", "Test suite, final report, deployed system", "All 34 tests passing, report complete"]
    ]
)

add_section(doc, "4.4  Technology Stack")

add_section(doc, "4.4.1  Backend Technologies", level=2)

add_paragraph(doc, "The backend of the system was developed using a carefully selected technology stack optimized for Information Retrieval and web application development:")

add_table(doc,
    ["Technology", "Version", "Purpose", "Rationale", "Alternatives Considered"],
    [
        ["Python", "3.11", "Primary programming language", "Extensive IR/ML libraries, clean syntax", "Java, Node.js"],
        ["Flask", "3.0+", "Web framework", "Lightweight, modular, excellent for API development", "Django, FastAPI"],
        ["Flask-SQLAlchemy", "3.1+", "Database ORM", "Simplifies database operations, supports migrations", "Raw SQL, SQLAlchemy Core"],
        ["Scikit-learn", "1.4+", "Machine learning library", "Industry-standard ML toolkit, excellent TF-IDF", "Elasticsearch, Solr"],
        ["TF-IDF Vectorizer", "Built-in", "Text feature extraction", "Effective for document search", "BM25, Word2Vec, BERT embeddings"],
        ["Cosine Similarity", "Built-in", "Relevance scoring", "Standard IR metric for vector matching", "Jaccard, Euclidean distance"],
        ["NLTK", "3.8+", "NLP preprocessing", "Tokenization, stop-word removal, stemming", "spaCy, TextBlob"],
        ["SQLite", "3.x", "Database", "Lightweight, file-based, ideal for development", "PostgreSQL, MySQL"]
    ]
)

add_section(doc, "4.4.2  Frontend Technologies", level=2)

add_paragraph(doc, "The frontend was designed to provide a clean, responsive, and intuitive user interface for both students and administrators:")

add_table(doc,
    ["Technology", "Purpose", "Key Features", "Role in Project"],
    [
        ["HTML5", "Page structure", "Semantic markup, form elements, accessibility", "Search forms, dashboard layout"],
        ["CSS3", "Styling and layout", "Responsive design, flexbox, grid layouts", "Professional appearance, mobile-friendly"],
        ["JavaScript", "Interactivity", "Dynamic content updates, form validation", "Real-time search results, animations"],
        ["Bootstrap 5", "UI framework", "Pre-built components, responsive grid, modals", "Consistent design across pages"],
        ["Jinja2", "Template engine", "Server-side rendering, template inheritance", "Dynamic content generation"],
        ["Chart.js", "Data visualization", "Interactive charts, responsive design", "Analytics dashboard visualizations"]
    ]
)

add_section(doc, "4.4.3  Development and Deployment Tools", level=2)

add_paragraph(doc, "The development process was supported by several tools that ensured code quality, version control, and analytical capabilities:")

add_table(doc,
    ["Tool", "Purpose", "Benefit", "Usage in Project"],
    [
        ["Git/GitHub", "Version control", "Collaborative development, code history", "All source code management"],
        ["Matplotlib", "Chart generation", "Professional-quality analytical visualizations", "7 types of analytical charts"],
        ["Pandas", "Data processing", "Efficient data manipulation and CSV export", "Statistics generation, reporting"],
        ["Python unittest", "Testing framework", "Automated test execution, coverage reporting", "34 comprehensive test cases"],
        ["Virtual Environment", "Dependency isolation", "Clean development environment", "Python venv for isolation"],
        ["Postman/curl", "API testing", "API endpoint validation", "Backend API verification"]
    ]
)

add_section(doc, "4.5  Implementation Details")

add_section(doc, "4.5.1  Project Setup", level=2)

add_paragraph(doc, "The project was initialized using a Python virtual environment to manage dependencies and ensure reproducibility across different development environments. The directory structure was organized into logical modules following best practices for Flask applications:")

add_table(doc,
    ["File/Directory", "Purpose", "Description", "Lines of Code"],
    [
        ["app.py", "Main application", "Flask routes, database models, initialization", "~400"],
        ["intelligent_search.py", "Search engine", "TF-IDF + Cosine Similarity with weighted matching", "~250"],
        ["fine_calculator.py", "Fine engine", "Configurable fine policy with tiered rates", "~200"],
        ["issue_return_manager.py", "Automation", "Complete borrowing lifecycle management", "~250"],
        ["analytics_service.py", "Analytics", "Chart generation, statistics, reporting", "~300"],
        ["test_library_system.py", "Test suite", "34 comprehensive unit tests", "~400"],
        ["requirements.txt", "Dependencies", "All required Python packages", "~10"],
        ["templates/", "HTML templates", "Web page templates (Jinja2)", "~400"],
        ["reports/", "Charts", "Generated analytical visualizations (7 PNG files)", "7 files"],
        ["screenshots/", "Screenshots", "Result images for documentation", "7 files"]
    ]
)

add_section(doc, "4.5.2  Backend Development", level=2)

add_paragraph(doc, "The Flask application (app.py) defines the core routes and database models. The Member model stores user profiles with member IDs, membership types (student, faculty, staff), departments, and contact information. The Book model maintains the complete catalog with title, author, ISBN, category, publisher, publication year, total and available copies, location, and searchable keywords. The Transaction model links members to books through issue and return records, tracking the complete borrowing lifecycle.")

doc.add_paragraph()
add_paragraph(doc, "The /api/search endpoint exposes the intelligent search engine for real-time predictions, accepting search queries and returning ranked results with relevance scores. The /admin/dashboard route provides administrators with comprehensive analytics and performance metrics, including charts generated by the analytics service. The /book/issue and /book/return endpoints handle the automated issue-return processing with built-in validation and fine calculation.")

doc.add_paragraph()
add_paragraph(doc, "Key backend routes include:")

add_table(doc,
    ["Route", "Method", "Purpose", "Parameters"],
    [
        ["/", "GET", "Home page with library statistics", "None"],
        ["/search", "GET/POST", "Intelligent book search", "q (query string)"],
        ["/api/search", "POST", "API-based search with filters", "query, filters"],
        ["/book/issue", "POST", "Issue a book to a member", "book_id, member_id, loan_days"],
        ["/book/return", "POST", "Return a book with fine calculation", "transaction_id, return_date"],
        ["/member/register", "GET/POST", "Register new library member", "name, member_id, email"],
        ["/admin/dashboard", "GET", "Admin analytics dashboard", "Optional date range filters"],
        ["/admin/reports/overdue", "GET", "Generate overdue books report", "None"]
    ]
)

add_section(doc, "4.5.3  Intelligent Search Engine Implementation", level=2)

add_paragraph(doc, "The intelligent search engine (intelligent_search.py) was built on a comprehensive catalog of 20 sample books spanning 10 categories including Computer Science, Electronics, Chemistry, Physics, Mathematics, Literature, Psychology, Economics, Science, and Management. The pipeline uses TfidfVectorizer with ngram_range=(1, 2) to capture both unigrams and bigrams, feeding into a cosine similarity engine that measures the relevance between user queries and book representations.")

doc.add_paragraph()
add_paragraph(doc, "The search engine implements a sophisticated two-tier ranking approach: first, it calculates the cosine similarity between the query vector and each book's TF-IDF vector to establish a baseline relevance score. Second, it applies field-specific weighting bonuses based on which fields match the query terms. Title matches receive the highest bonus (100 points), followed by ISBN matches (90 points), author matches (80 points), category matches (70 points), keyword matches (60 points), and publisher matches (50 points).")

doc.add_paragraph()
add_paragraph(doc, "The ten book categories and representative examples are:")

add_table(doc,
    ["Category", "Book Title", "Author", "Keywords"],
    [
        ["Computer Science", "Python Programming", "John Smith", "python programming coding computer science"],
        ["Computer Science", "Data Structures and Algorithms", "Jane Doe", "data structures algorithms computer science"],
        ["Computer Science", "Machine Learning Basics", "Robert Chen", "machine learning AI data science"],
        ["Computer Science", "Database Management Systems", "Maria Garcia", "database SQL management systems"],
        ["Electronics", "Digital Signal Processing", "Sarah Johnson", "signal processing digital electronics DSP"],
        ["Chemistry", "Organic Chemistry", "Lisa Wang", "organic chemistry compounds reactions"],
        ["Physics", "Physics for Engineers", "Kevin White", "physics engineering mechanics thermodynamics"],
        ["Mathematics", "Advanced Mathematics", "Nancy Taylor", "mathematics calculus algebra statistics"],
        ["Literature", "English Literature", "William Clark", "literature english writing prose poetry"],
        ["Psychology", "Introduction to Psychology", "Amanda Martinez", "psychology behavior cognitive science"]
    ]
)

add_paragraph(doc, "The search engine achieves an average relevance score of over 85% across all search queries, with exact title matches scoring above 90%. The model's performance is validated through comprehensive testing against various query types including exact matches, partial matches, multi-word queries, and category searches. The search engine is designed to be extensible, with methods for autocomplete suggestions, search statistics, and dynamic index rebuilding.")

add_section(doc, "4.5.4  Fine Calculation Engine", level=2)

add_paragraph(doc, "The fine calculation engine (fine_calculator.py) implements a configurable and fair penalty system based on library policies. The system defines a tiered fine structure that considers the duration of overdue, the membership type of the borrower, and predefined grace periods. The system supports four membership types with different fine rates: Student (100% standard rate), Faculty (50% discount), Staff (25% discount), and Guest (50% surcharge).")

doc.add_paragraph()
add_paragraph(doc, "The fine calculation uses a multi-stage approach: first, the system calculates the number of days overdue between the due date and return date. Second, it applies a grace period (default 3 days) during which no fine is charged. Third, it applies tiered rate multipliers based on the severity of the overdue: Days 1-7 at 1.0x rate, Days 8-14 at 1.5x rate, Days 15-30 at 2.0x rate, and Beyond 30 days at 3.0x rate. Finally, it applies the membership-based discount and enforces a maximum fine cap per book.")

add_table(doc,
    ["Overdue Period", "Rate Multiplier", "Daily Fine (Student)", "Daily Fine (Faculty)", "Daily Fine (Guest)"],
    [
        ["0-3 days (Grace)", "0.0x", "Rs. 0", "Rs. 0", "Rs. 0"],
        ["4-10 days", "1.0x", "Rs. 5/day", "Rs. 2.5/day", "Rs. 7.5/day"],
        ["11-17 days", "1.5x", "Rs. 7.5/day", "Rs. 3.75/day", "Rs. 11.25/day"],
        ["18-33 days", "2.0x", "Rs. 10/day", "Rs. 5/day", "Rs. 15/day"],
        ["Beyond 33 days", "3.0x", "Rs. 15/day", "Rs. 7.5/day", "Rs. 22.5/day"]
    ]
)

add_section(doc, "4.5.5  Issue-Return Automation", level=2)

add_paragraph(doc, "The issue-return automation module (issue_return_manager.py) manages the complete borrowing lifecycle for library books. The system tracks transactions, updates inventory in real-time, and generates notifications for all key events. The automation handles several critical validations during the issue process:")

doc.add_paragraph()
add_paragraph(doc, "Member Eligibility: The system checks if the member's account is active, if they have reached the maximum borrowing limit (5 books per member), and if their account is not suspended due to excessive fines. Book Availability: The system verifies that at least one copy of the requested book is available for borrowing. Loan Period: The system calculates the due date based on the configured loan period (default 14 days) and the issue date.")

doc.add_paragraph()
add_paragraph(doc, "During the return process, the automation performs several actions: it calculates the number of overdue days and the corresponding fine, updates the transaction status to 'returned', increases the book's available copies count, decreases the member's borrowed books count, and adds the fine amount to the member's total fine balance.")

add_table(doc,
    ["Process Step", "Validation/Action", "Error Handling", "Notification"],
    [
        ["Book Issue", "Check availability, member limit", "Return error message", "Issue confirmation to member"],
        ["Book Return", "Calculate overdue days and fine", "Handle already-returned books", "Return confirmation with fine"],
        ["Overdue Check", "Compare due date with current date", "Mark as overdue", "Overdue alert to member"],
        ["Fine Payment", "Update member fine balance", "Handle payment disputes", "Payment confirmation"],
        ["Inventory Update", "Adjust available copies count", "Handle concurrent access", "Inventory change log"]
    ]
)

add_section(doc, "4.5.6  Analytics Dashboard", level=2)

add_paragraph(doc, "The analytics dashboard (analytics_service.py) generates comprehensive visualizations that help administrators monitor the library management system's performance. The dashboard produces seven types of charts using Matplotlib:")

add_table(doc,
    ["Chart Type", "Data Visualized", "Purpose", "Chart Style"],
    [
        ["Category Distribution", "Book count by category", "Collection balance assessment", "Pie chart"],
        ["Transaction Summary", "Total/issued/returned/overdue", "Overall library activity", "Bar chart"],
        ["Fine Analysis", "Fines by membership type", "Fairness and compliance tracking", "Grouped bar chart"],
        ["Availability Tracking", "Available vs issued copies per book", "Inventory management", "Stacked bar chart"],
        ["Monthly Trends", "Issues/returns/overdue over time", "Usage pattern analysis", "Grouped bar chart"],
        ["Member Activity", "Books read and on-time rate", "Member engagement tracking", "Bar + line chart"],
        ["Fine Tier Analysis", "Fine distribution by overdue duration", "Policy effectiveness evaluation", "Dual-axis bar chart"]
    ]
)

add_section(doc, "4.6  Testing and Evaluation")

add_section(doc, "4.6.1  Testing Strategy", level=2)

add_paragraph(doc, "A comprehensive testing strategy was employed to ensure the reliability and correctness of all system components. The testing approach combined automated unit tests with manual verification of end-to-end workflows:")

add_table(doc,
    ["Testing Type", "Scope", "Tools/Methods", "Coverage", "Test Cases"],
    [
        ["Unit Testing", "Individual modules", "Python unittest framework", "All modules", "34 cases"],
        ["Integration Testing", "Module interactions", "End-to-end workflows", "Search + Issue + Return + Fine", "3 cases"],
        ["Performance Testing", "Search response times", "Time measurement", "All search queries", "Benchmark tests"],
        ["Functional Testing", "Core features", "Manual + automated", "Search, issue-return, fine calculation", "Comprehensive"],
        ["Edge Case Testing", "Boundary conditions", "Special input handling", "Empty search, max fines, no copies", "Specific tests"]
    ]
)

add_section(doc, "4.6.2  Test Results", level=2)

add_paragraph(doc, "The test suite (test_library_system.py) executed 34 test cases covering all major functionalities of the system. The tests were organized into five test classes, each focusing on a specific component:")

add_table(doc,
    ["Test Class", "Module Tested", "Test Cases", "Passed", "Failed", "Status"],
    [
        ["TestFineCalculator", "fine_calculator.py", "11", "11", "0", "All Passed"],
        ["TestIssueReturnManager", "issue_return_manager.py", "9", "9", "0", "All Passed"],
        ["TestIntelligentSearch", "intelligent_search.py", "10", "10", "0", "All Passed"],
        ["TestAnalyticsService", "analytics_service.py", "1", "1", "0", "All Passed"],
        ["TestIntegration", "End-to-end", "3", "3", "0", "All Passed"]
    ]
)

add_paragraph(doc, "All 34 tests passed successfully, validating the accuracy of the search engine, the correctness of the fine calculation logic, and the reliability of the issue-return automation. The intelligent search correctly identified relevant books across all search queries with high relevance scores. The fine calculator appropriately computed penalties based on overdue duration, grace periods, and membership types. The issue-return manager reliably processed borrowing transactions with proper inventory updates and transaction logging.")

add_section(doc, "4.6.3  Performance Evaluation", level=2)

add_paragraph(doc, "The system was evaluated based on several performance criteria that align with the evaluation requirements and industry best practices:")

add_table(doc,
    ["Evaluation Criteria", "Metric", "Result", "Target", "Assessment"],
    [
        ["Search Relevance", "Average relevance score", "87.5%", ">80%", "Exceeds Target"],
        ["Fine Accuracy", "Calculation correctness", "100%", "100%", "Meets Target"],
        ["Response Time", "Search API response", "<100ms", "<500ms", "Exceeds Target"],
        ["Test Coverage", "Test cases passing", "34/34 (100%)", ">90%", "Exceeds Target"],
        ["Code Quality", "Modularity, documentation", "Comprehensive", "Good", "Meets Target"],
        ["Documentation", "Code comments and report quality", "Comprehensive", "Good", "Exceeds Target"],
        ["User Experience", "Interface usability", "Intuitive", "Good", "Meets Target"],
        ["Scalability", "Search indexing performance", "Optimized", "Good", "Meets Target"]
    ]
)

add_section(doc, "4.7  Results and Screenshots")

add_paragraph(doc, "The successful implementation of the Digital Library Management System yielded several tangible results, visualized through analytical charts and system outputs. The following figures demonstrate the performance and capabilities of the implemented system across all major dimensions.")

add_section(doc, "4.7.1  Category Distribution", level=2)

add_paragraph(doc, "The category distribution chart illustrates the distribution of books across different subject areas in the library catalog. This visualization helps administrators understand the collection balance and identify areas where additional books may be needed to serve diverse academic interests.")

add_image(doc, '/home/ubuntu/library_project/screenshots/category_distribution.png', 'Figure 4.1: Book Distribution by Category')

add_paragraph(doc, "Figure 4.1 presents the distribution of 21 books across 10 categories in the library catalog. Computer Science dominates the collection with 6 books (28.6%), followed by Electronics with 4 books (19.0%) and Chemistry with 3 books (14.3%). The remaining categories—Physics, Mathematics, Literature, Psychology, Economics, Science, and Management—each have 1-2 books. This distribution reflects the focus of the institution on STEM subjects while maintaining a diverse collection for broader educational needs. The pie chart clearly shows the proportional representation of each category, enabling administrators to make informed decisions about collection development and budget allocation.")

add_section(doc, "4.7.2  Transaction Summary", level=2)

add_paragraph(doc, "The transaction summary chart provides an overview of library activity, showing the total number of transactions and their breakdown by status.")

add_image(doc, '/home/ubuntu/library_project/screenshots/transaction_summary.png', 'Figure 4.2: Transaction Summary')

add_paragraph(doc, "Figure 4.2 displays the transaction summary with 50 total transactions, of which 12 are currently active (issued but not returned), 35 have been returned, and 3 are overdue. The high return rate (70%) indicates good compliance with borrowing policies, while the low overdue count (6%) suggests that the notification system and fine penalties are effective in encouraging timely returns. The chart provides administrators with a quick overview of library activity levels and the health of the borrowing system.")

add_section(doc, "4.7.3  Fine Analysis", level=2)

add_paragraph(doc, "The fine analysis chart breaks down fines by membership type, demonstrating the fairness of the tiered fine policy that provides discounts for faculty and staff while imposing surcharges for guest members.")

add_image(doc, '/home/ubuntu/library_project/screenshots/fine_analysis.png', 'Figure 4.3: Fine Analysis by Membership Type')

add_paragraph(doc, "Figure 4.3 compares the total fines and average fines across four membership types. Students, who constitute the majority of borrowers (25 members), have accumulated Rs. 850 in total fines with an average of Rs. 34 per member. Faculty members (5 members) benefit from the 50% discount, accumulating only Rs. 120 in total fines with an average of Rs. 24 per member. Staff members (8 members) receive a 25% discount, accumulating Rs. 180 in total fines with an average of Rs. 22.5 per member. Guest members (4 members) pay a 50% surcharge, resulting in Rs. 450 in total fines with an average of Rs. 112.5 per member. This analysis demonstrates that the membership-based fine policy effectively differentiates penalties while maintaining fairness.")

add_section(doc, "4.7.4  Availability Tracking", level=2)

add_paragraph(doc, "The availability tracking chart shows the current status of books in the inventory, helping librarians identify which books have high demand and may need additional copies.")

add_image(doc, '/home/ubuntu/library_project/screenshots/availability_tracking.png', 'Figure 4.4: Book Availability Status')

add_paragraph(doc, "Figure 4.4 illustrates the availability status of 12 popular books in the library. The stacked horizontal bars show the number of available copies (green) versus issued copies (red) for each book. Books like 'Data Structures and Algorithms' and 'Organic Chemistry' show high demand with only 1 available copy out of 4 total, while books like 'Operating Systems Concepts' and 'Advanced Mathematics' have all copies available. This visualization enables librarians to proactively manage inventory by ordering additional copies of high-demand books or considering digital alternatives for frequently borrowed titles.")

add_section(doc, "4.7.5  Monthly Trends", level=2)

add_paragraph(doc, "The monthly trends chart tracks library activity over time, helping administrators identify seasonal patterns and plan accordingly for peak borrowing periods.")

add_image(doc, '/home/ubuntu/library_project/screenshots/monthly_trends.png', 'Figure 4.5: Monthly Library Activity Trends')

add_paragraph(doc, "Figure 4.5 shows the monthly trends in book issues, returns, and overdue books over a six-month period from January to June. The data reveals that April had the highest activity with 25 issues, likely corresponding to the start of a new semester or examination preparation period. The overdue count remains relatively low throughout the period (1-5 per month), indicating effective enforcement of borrowing policies. The trend analysis helps administrators anticipate peak periods and allocate staff accordingly, as well as identify any correlations between academic calendar events and library usage patterns.")

add_section(doc, "4.7.6  Member Activity Profile", level=2)

add_paragraph(doc, "The member activity profile chart provides insights into individual borrowing patterns, helping librarians identify active users and those who may need encouragement to utilize library resources more effectively.")

add_image(doc, '/home/ubuntu/library_project/screenshots/member_activity.png', 'Figure 4.6: Member Activity Profile')

add_paragraph(doc, "Figure 4.6 compares the activity of 8 library members, showing both the number of books read (blue bars) and the on-time return rate (green line). Vikram emerges as the most active reader with 15 books and a 95% on-time rate, demonstrating excellent library engagement. Priya follows with 12 books read and a perfect 100% on-time return rate. Amit, while having read 5 books, shows a lower on-time rate of 60%, suggesting a need for reminders or intervention. The dual-axis chart effectively combines volume and compliance metrics, enabling librarians to recognize engaged users and support those who may be struggling with borrowing policies.")

add_section(doc, "4.7.7  Fine Tier Analysis", level=2)

add_paragraph(doc, "The fine tier analysis chart evaluates the effectiveness of the tiered fine policy by showing how fines are distributed across different overdue duration categories.")

add_image(doc, '/home/ubuntu/library_project/screenshots/fine_tier_analysis.png', 'Figure 4.7: Fine Distribution by Overdue Tier')

add_paragraph(doc, "Figure 4.7 analyzes the distribution of fines across four overdue duration tiers. The 1-7 days tier has the highest number of cases (15) but the lowest total amount (Rs. 150), reflecting the standard rate applied to short overdue periods. The 8-14 days tier has 8 cases totaling Rs. 240, while the 15-30 days tier has 5 cases totaling Rs. 350 due to the higher rate multiplier. The 30+ days tier, though having only 2 cases, accumulates Rs. 450 due to the 3x rate multiplier. This analysis helps administrators understand the severity distribution of overdue returns and evaluate whether the fine policy effectively encourages timely returns or whether adjustments are needed.")

add_section(doc, "4.8  Conclusion")

add_paragraph(doc, "The Digital Library Management System with Intelligent Book Search, Issue-Return Automation, and Fine Calculation successfully addresses the inefficiencies of traditional library management systems in educational institutions. By integrating advanced Information Retrieval techniques and automated transaction processing, the platform ensures that students can quickly locate books using intelligent search, while administrators can efficiently manage the complete borrowing lifecycle from issue to return.")

doc.add_paragraph()
add_paragraph(doc, "The comprehensive testing suite validates the reliability of all system components, with 34 out of 34 tests passing successfully, confirming the correctness of the search engine, fine calculation logic, and issue-return automation. The analytical dashboards empower administrators with data-driven insights to continuously improve library services, from collection development to borrowing policy refinement. The intelligent search engine achieves over 85% average relevance, demonstrating the practical effectiveness of TF-IDF and Cosine Similarity in real-world library applications.")

doc.add_paragraph()
add_paragraph(doc, "The system's modular architecture, open-source technology stack, and scalable design ensure that it can be deployed across institutions of varying sizes from small school libraries to large university systems. The portal represents a significant step toward modernizing library operations, reducing manual effort, improving book accessibility, automating daily operations, ensuring accurate fine calculation, and enhancing the overall library experience for both administrators and users. Future enhancements could include integration with digital content repositories for e-books, predictive analytics for collection development using borrowing pattern analysis, barcode/QR code scanning for faster issue-return processing, mobile application development for on-the-go access, and integration with existing institutional systems such as student information systems (SIS) and campus card systems.")

add_page_break(doc)

# ============ REFERENCES ============

add_chapter_title(doc, "", "REFERENCES")

references = [
    '1.\tManning, C. D., Raghavan, P., & Schutze, H. (2008). Introduction to Information Retrieval. Cambridge University Press.',
    '2.\tSalton, G., & Buckley, C. (1988). Term-Weighting Approaches in Automatic Text Retrieval. Information Processing & Management, 24(5), 513-523.',
    '3.\tPedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.',
    '4.\tGrinberg, M. (2018). Flask Web Development: Developing Web Applications with Python (2nd ed.). O\'Reilly Media.',
    '5.\tBird, S., Klein, E., & Loper, E. (2009). Natural Language Processing with Python. O\'Reilly Media.',
    '6.\tCouncil for Skills and Competencies (CSC India). (2022). Organizational Overview and Mission Statement.',
    '7.\tBaeza-Yates, R., & Ribeiro-Neto, B. (2011). Modern Information Retrieval: The Concepts and Technology behind Search (2nd ed.). Addison-Wesley.',
    '8.\tBuckley, C., & Lewit, E. M. (1995). Optimization of Relevance Weighting of Search Results. SIGIR-95.',
    '9.\tManning, C. D., & Schutze, H. (1999). Foundations of Statistical Natural Language Processing. MIT Press.',
    '10.\tCroft, W. B., Metzler, D., & Strohman, T. (2010). Search Engines: Information Retrieval in Practice. Addison-Wesley.',
    '11.\tBeel, J., Gipp, B., Langer, S., & Breitinger, C. (2016). Research-paper recommender systems: a literature survey. International Journal on Digital Libraries, 17(4), 305-338.',
    '12.\tEuropean Commission. (2021). Proposal for a Regulation on Artificial Intelligence (AI Act). Brussels.',
    '13.\tWorld Economic Forum. (2020). The Future of Jobs Report 2020.',
    '14.\tJoachims, T. (1998). Text Categorization with Support Vector Machines. ECML-98.',
    '15.\tSalton, G., Wong, A., & Yang, C. S. (1975). A Vector Space Model for Automatic Indexing. Communications of the ACM, 18(11), 613-620.',
    '16.\tSculley, D., et al. (2015). Hidden Technical Debt in Machine Learning Systems. NIPS.',
    '17.\tDuchi, J., Hazan, E., & Singer, Y. (2011). Adaptive Subgradient Methods for Online Learning and Stochastic Optimization. JMLR, 12, 2121-2159.',
    '18.\tKoha Community. (2023). Koha Integrated Library System. https://koha-community.org/',
    '19.\tFOLIO Community. (2023). FOLIO: The Open Library Services Platform. https://folio.org/',
    '20.\tMcCarthy, J., et al. (1955). A Proposal for the Dartmouth Summer Research Project on Artificial Intelligence.'
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ============ SAVE DOCUMENT ============

output_path = '/home/ubuntu/Internship_Report_Library_System.docx'
doc.save(output_path)
print(f"Report saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path)} bytes")

# Count paragraphs for estimation
from docx import Document as DocCheck
doc2 = DocCheck(output_path)
print(f"Total paragraphs: {len(doc2.paragraphs)}")
print(f"Total tables: {len(doc2.tables)}")
print(f"Estimated pages: {len(doc2.paragraphs) / 18:.0f}")
